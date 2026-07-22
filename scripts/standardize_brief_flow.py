"""Replace the MARACOOS pathway strip with four equal-width stages."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


PALE = "EAF5F6"
TEAL = RGBColor.from_string("007C89")
GOLD = RGBColor.from_string("E8A43A")
GRAY = RGBColor.from_string("58666D")
TOTAL_WIDTH_DXA = 10627
CELL_WIDTHS_DXA = (2657, 2657, 2657, 2656)


def set_cell_width(cell, width: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    cell_width = properties.find(qn("w:tcW"))
    if cell_width is None:
        cell_width = OxmlElement("w:tcW")
        properties.append(cell_width)
    cell_width.set(qn("w:w"), str(width))
    cell_width.set(qn("w:type"), "dxa")


def style_cell(cell, title: str, subtitle: str, show_arrow: bool) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), PALE)

    margins = OxmlElement("w:tcMar")
    for edge, value in (("top", 110), ("start", 95), ("bottom", 100), ("end", 95)):
        item = OxmlElement(f"w:{edge}")
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")
        margins.append(item)
    properties.append(margins)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    cell.text = ""
    title_paragraph = cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(2)
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(8.4)
    title_run.font.color.rgb = TEAL
    if show_arrow:
        arrow_run = title_paragraph.add_run("  →")
        arrow_run.bold = True
        arrow_run.font.name = "Arial"
        arrow_run.font.size = Pt(9)
        arrow_run.font.color.rgb = GOLD

    subtitle_paragraph = cell.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.paragraph_format.space_after = Pt(0)
    subtitle_paragraph.paragraph_format.line_spacing = 0.95
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(7.2)
    subtitle_run.font.color.rgb = GRAY


def standardize(source: Path, destination: Path) -> None:
    document = Document(source)
    host_table = next(
        table
        for table in document.tables
        if any("OBSERVE" in cell.text and "PREDICT" in " ".join(c.text for c in row.cells)
               for row in table.rows for cell in row.cells)
    )
    flow_row = next(row for row in host_table.rows if any("OBSERVE" in cell.text for cell in row.cells))

    flow_table = document.add_table(rows=1, cols=4)
    flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    flow_table.autofit = False
    table_properties = flow_table._tbl.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(TOTAL_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    table_properties.append(borders)

    grid = flow_table._tbl.tblGrid
    for grid_column, width in zip(grid.gridCol_lst, CELL_WIDTHS_DXA):
        grid_column.set(qn("w:w"), str(width))

    stages = (
        ("OBSERVE", "radar • buoys • sensors"),
        ("CONNECT", "quality-controlled data"),
        ("PREDICT", "models • forecasts"),
        ("DECIDE", "search • navigate • prepare • operate"),
    )
    for index, (cell, (title, subtitle), width) in enumerate(
        zip(flow_table.rows[0].cells, stages, CELL_WIDTHS_DXA)
    ):
        set_cell_width(cell, width)
        style_cell(cell, title, subtitle, show_arrow=index < len(stages) - 1)

    host_table._tbl.addnext(flow_table._tbl)
    host_table._tbl.remove(flow_row._tr)
    document.save(destination)


if __name__ == "__main__":
    standardize(Path(sys.argv[1]), Path(sys.argv[2]))
